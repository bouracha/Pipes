import pytesseract
import cv2
import os
from docx import Document
from PIL import Image
import io

# Function to extract text from images using OCR
def extract_text_from_image(image_bytes):
    image = Image.open(io.BytesIO(image_bytes))
    text = pytesseract.image_to_string(image)
    return text.strip()  # Trim extra whitespace

# Function to extract text from DOCX, including tables & OCR for embedded images
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    extracted_text = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():  # Avoid empty lines
            extracted_text.append(para.text.strip())

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_text:  # Avoid empty rows
                extracted_text.append(row_text)

    # Extract OCR text from embedded images
    for rel in doc.part.rels:
        if "image" in doc.part.rels[rel].target_ref:
            image_data = doc.part.rels[rel].target_part.blob
            extracted_text.append("[Image OCR] " + extract_text_from_image(image_data))

    return "\n".join(extracted_text)  # Join everything with newlines

# Path to your SOP document
docx_path = "sop.docx"
extracted_text = extract_text_from_docx(docx_path)

# Save extracted text to a file
with open("extracted_text.txt", "w", encoding="utf-8") as f:
    f.write(extracted_text)

print("OCR extraction complete. Check extracted_text.txt")
